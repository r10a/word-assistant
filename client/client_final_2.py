from docx import Document
import json
import websocket
import os


def on_message(ws, message_from_cloud):
    print('Message received')
    print(message_from_cloud)
    if message_from_cloud[0:3] != 'Ack':
        message_clientID, message_filename, message_command, message_value, message_count = parse_message(
            message_from_cloud)
        if message_clientID == '2':
            document_name = 'outputs/' + message_filename + '.docx'
            doc = DocumentWriter(docname=document_name)
            # Create a file
            if message_command == 'create':
                try:
                    doc.add_paragraph(message_value)
                    doc.save_document()
                    ws.send("Ack,1,0")
                except:
                    ws.send("Ack,1,1")

            # Edit an existing File
            if message_command == 'edit':
                try:
                    doc.add_paragraph(message_value)
                    doc.save_document()
                    ws.send("Ack,1,0")
                except:
                    ws.send("Ack,1,1")

            # Delete a file
            if message_command == 'delete':
                try:
                    os.remove(document_name)
                    ws.send("Ack,1,0")
                except OSError:
                    ws.send("Ack,1,1")
        else:
            print('Not Client 2 message')


def on_error(ws, error):
    print(error)


def on_close(ws):
    ws.close()
    print("### closed ###")


def on_open(ws):
    ws.send("Hello,2")


# TODO: Add more functionality
class DocumentWriter:
    def __init__(self, docname):
        self._current_para = None
        self.error_code = 0
        self._document = Document()
        self._docname = docname

    def add_text(self, text):
        self._current_para.add_run(text=text)

    def add_heading(self, text):
        self._current_para = self._document.add_heading(text=text)

    def add_paragraph(self, text):
        self._current_para = self._document.add_paragraph(text=text)

    def save_document(self):
        self._document.save(self._docname)


def parse_message(message_from_gce):
    json_data = json.loads(message_from_gce)
    message_clientID = str(json_data['clientID'])
    message_filename = str(json_data['filename'])
    message_command = str(json_data['command'])
    message_value = str(json_data['text'])
    message_count = str(json_data['counter'])
    return message_clientID, message_filename, message_command, message_value, message_count


# TODO: Configure Websocket connection to Cloud server
from websocket import create_connection

websocket.enableTrace(True)
ws = websocket.WebSocketApp("ws://127.0.0.1:8080/websocket",
                            on_message=on_message,
                            on_error=on_error,
                            on_close=on_close)
ws.on_open = on_open
ws.run_forever()