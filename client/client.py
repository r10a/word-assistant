from docx import Document
import json
import logging
import redis
import os
from pywinauto.application import Application

from uuid import getnode

REDIS_URL = "redis://h:p4a49c92f2f92f61555110cca953dd9b8fc55fe9736e8aa3d277ea93fa4abb0c0@ec2-54-158-0-180.compute-1.amazonaws.com:62409"
REDIS_SERVER = 'server'
REDIS_CLIENT = 'client'

redis = redis.from_url(REDIS_URL)


# TODO: Add more functionality
class DocumentWriter:

    def __init__(self, docname="generated.docx"):
        self._current_para = None
        self.error_code = 0
        #self._document = Document()
        self._document = Application().start("C:\\Program Files\\Microsoft Office\\Office16\\WINWORD.EXE", timeout=50)
        self._new_window = self._document.Window_(top_level_only=True, active_only=True, class_name='OpusApp')
        self._docname = docname

    def get_docname(self):
        return self._docname

    def add_text(self, text):
        if self._current_para is None:
            self.add_paragraph(text)
        else:
            self._current_para.add_run(text=text)

    def add_heading(self, text):
        self._current_para = self._document.add_heading(text=text)

    def add_paragraph(self, text):
        self._current_para = self._document.add_paragraph(text=text)

    # def save_document(self):
    #     self._document.save(self._docname)


class CommandHandler:

    def __init__(self):
        self.doc = None
        self.new_window = None
        self.document_open = False
        self.pubsub = redis.pubsub()
        self.pubsub.subscribe(REDIS_SERVER)

    def open_doc(self):
        self.doc = Application(backend='uia').start("C:\\Program Files\\Microsoft Office\\Office16\\WINWORD.EXE", timeout=50)
        self.new_window = self.doc.window(top_level_only=True, active_only=True, class_name='OpusApp')
        self.document_open = True
        self.respond(True, "")

    def __iter_data(self):
        for message in self.pubsub.listen():
            data = message.get('data')
            if message['type'] == 'message':
                yield data

    def run(self):
        """Listens for new commands in Redis, and executes them."""
        for data in self.__iter_data():
            self.on_message(data)

    def respond(self, completed, parameters):
        ack = {
            "id": getnode(),
            "completed": completed,
            "parameters": parameters
        }
        print(ack)
        redis.publish(REDIS_CLIENT, json.dumps(ack))

    def save_doc(self):
        self.new_window.type_keys(r'%FAO{ENTER}')


    def on_message(self, message):
        print("command", message)
        command, parameters = self.parse_message(message)

        if command == 'open':
            try:
                print("opening document")
                self.open_doc()
                self.respond(True, "Opened")
            except Exception as e:
                logging.error('Failed to open document: ' + str(e))
                self.respond(False, "Not Opened")

        if command == 'type':
            try:
                if not self.document_open:
                    self.open_doc()
                print("adding", parameters)
                self.new_window.type_keys(parameters.replace(" ", "{SPACE}") + ".{SPACE}")
                self.respond(True, parameters)
            except Exception as e:
                logging.error('Failed to add to document: ' + str(e))
                self.respond(False, parameters)

        if command == 'delete':  # TODO: Implement
            try:
                print("deleting", parameters)
                os.remove(self.doc.get_docname())
                self.respond(True, "Deleted")
            except Exception as e:
                logging.error('Failed to delete document: ' + str(e))
                self.respond(False, "Not Deleted")

        if command == 'save':
            try:
                print("saving")
                self.save_doc()
                self.respond(True, "Saved")
            except Exception as e:
                logging.error('Failed to save document: ' + str(e))
                self.respond(False, "Not saved")

        if command == 'close':
            try:
                print("closing")
                self.new_window.type_keys(r'%{F4}')
                self.respond(True, "Closed")
            except Exception as e:
                logging.error('Failed to save document: ' + str(e))
                self.respond(False, "Not Closed")

    def parse_message(self, command):
        print("command", command)
        json_data = json.loads(command)
        query = str(json_data['queryText'])
        parameters = json_data['parameters']['text'] if json_data['parameters'] else None
        command = query.split()[0]
        return command, parameters


ch = CommandHandler()
ch.run()